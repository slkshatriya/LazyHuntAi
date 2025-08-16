import streamlit as st
import tempfile
import re
import requests
import spacy
import io
from bs4 import BeautifulSoup
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from docx import Document

# -------------------- NLP Model --------------------
@st.cache_resource
def get_nlp_model():
    # Model will already be installed via requirements.txt
    return spacy.load("en_core_web_sm")

nlp = get_nlp_model()

# -------------------- Resume PDF Builder --------------------
def build_resume(data, file_path):
    doc = SimpleDocTemplate(file_path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Header', fontSize=18, leading=22, spaceAfter=10, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='SubHeader', fontSize=13, leading=16, spaceAfter=6))
    styles.add(ParagraphStyle(name='List', fontSize=10.5, leading=14, spaceAfter=6))
    contact_info_style = ParagraphStyle(name='ContactInfo', parent=styles['Normal'], alignment=TA_CENTER)

    flow = []
    flow.append(Paragraph(data['name'], styles['Header']))
    flow.append(Paragraph(f"{data['location']} | {data['phone']} | {data['email']}", contact_info_style))
    flow.append(Spacer(1, 12))

    flow.append(Paragraph("Professional Summary", styles['SubHeader']))
    flow.append(Paragraph(data['summary'], styles['List']))

    flow.append(Paragraph("Experience", styles['SubHeader']))
    for exp in data['experience']:
        flow.append(Paragraph(f"{exp['role']} at {exp['company']} ({exp['duration']})", styles['List']))
        for line in exp['tasks'].split('\n'):
            flow.append(Paragraph(f"• {line.strip()}", styles['List']))

    flow.append(Paragraph("Education", styles['SubHeader']))
    for edu in data['education']:
        flow.append(Paragraph(f"{edu['degree']}, {edu['institution']} ({edu['duration']})", styles['List']))

    flow.append(Paragraph("Skills", styles['SubHeader']))
    flow.append(Paragraph(", ".join(data['skills']), styles['List']))

    doc.build(flow)

# -------------------- Job Description Extractor --------------------
def extract_job_desc(url: str) -> str:
    try:
        res = requests.get(url, timeout=10)
        soup = BeautifulSoup(res.content, 'html.parser')
        for cls in ['description', 'jobDescription', 'job-desc', 'job-description']:
            jd_div = soup.find(class_=cls)
            if jd_div:
                return jd_div.get_text(" ", strip=True)
        return soup.body.get_text(" ", strip=True)[:5000]
    except Exception as e:
        st.error(f"Error extracting job description: {e}")
        return ""

# -------------------- Skill Extractor --------------------
def extract_skills(text):
    keywords = {
        'Python', 'PowerShell', 'Azure', 'AWS', 'GCP', 'Kubernetes', 'MLflow', 'Kubeflow', 'RAG',
        'LLM', 'Langchain', 'Mistral', 'Llama', 'GPT-4', 'TensorFlow', 'PyTorch', 'Django',
        'Ansible', 'StackStorm', 'ADO', 'GitHub', 'Bitbucket', 'Vector DB', 'Databricks'
    }
    found = set()
    doc = nlp(text)
    for token in doc:
        if token.text in keywords:
            found.add(token.text)
    matches = re.findall(r'\b[A-Z][a-zA-Z\d\-\+\.\#]{2,}\b', text)
    found.update([m for m in matches if m in keywords])
    return sorted(found)

# -------------------- Resume .docx Reader --------------------
def read_docx(file) -> tuple[Document, list]:
    buffer = io.BytesIO(file.read())
    doc = Document(buffer)
    return doc, [p.text for p in doc.paragraphs]

# -------------------- Update Skills in Resume --------------------
def update_skills_only(doc: Document, lines: list, new_skills: list):
    header = "skills"
    idx = None
    for i, line in enumerate(lines):
        if header in line.lower():
            idx = i
            break
    if idx is not None:
        nexti = idx + 1
        existing = set()
        if nexti < len(lines):
            existing.update({s.strip().lower() for s in re.split(r',|/|;|\n', lines[nexti]) if s.strip()})
        merged = existing.union([s.lower() for s in new_skills])
        doc.paragraphs[nexti].text = ", ".join(sorted({s.title() for s in merged}))

# -------------------- Streamlit UI --------------------
st.title("💫 LazyHuntAi: Your Resume Assistant")
mode = st.radio("Choose an Option:", ["Build Resume from Scratch", "Update Resume Skills from Job URL"])

if mode == "Build Resume from Scratch":
    st.header("📄 Create Your Resume")
    name = st.text_input("Full Name")
    location = st.text_input("Location")
    phone = st.text_input("Phone Number")
    email = st.text_input("Email Address")
    summary = st.text_area("Professional Summary")

    experience = []
    exp_count = st.number_input("Number of Experiences", 1, 5, 1)
    for i in range(int(exp_count)):
        with st.expander(f"Experience {i+1}"):
            exp_company = st.text_input("Company", key=f"exp_company_{i}")
            exp_location = st.text_input("Location", key=f"exp_loc_{i}")
            exp_role = st.text_input("Role", key=f"exp_role_{i}")
            exp_duration = st.text_input("Duration", key=f"exp_dur_{i}")
            exp_tasks = st.text_area("Responsibilities (one per line)", key=f"exp_tasks_{i}")
            experience.append({
                'company': exp_company,
                'location': exp_location,
                'role': exp_role,
                'duration': exp_duration,
                'tasks': exp_tasks
            })

    education = []
    edu_count = st.number_input("Number of Education Entries", 1, 5, 1)
    for i in range(int(edu_count)):
        with st.expander(f"Education {i+1}"):
            degree = st.text_input("Degree", key=f"degree_{i}")
            institution = st.text_input("Institution", key=f"institution_{i}")
            duration = st.text_input("Duration", key=f"edu_dur_{i}")
            education.append({
                'degree': degree,
                'institution': institution,
                'duration': duration
            })

    skills = st.text_area("Enter skills (comma-separated)")
    if st.button("Generate Resume PDF"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            build_resume({
                'name': name,
                'location': location,
                'phone': phone,
                'email': email,
                'summary': summary,
                'experience': experience,
                'education': education,
                'skills': [s.strip() for s in skills.split(",") if s.strip()]
            }, tmp.name)
            st.success("✅ Resume generated successfully!")
            with open(tmp.name, "rb") as f:
                st.download_button("⬇️ Download Resume PDF", f, file_name="resume.pdf")

else:
    st.header("🔁 Update Skills in Existing Resume")
    job_url = st.text_input("Paste Job Post URL")
    uploaded_file = st.file_uploader("Upload Resume (.docx)", type=["docx"])
    if st.button("Update Skills from JD") and job_url and uploaded_file:
        jd_text = extract_job_desc(job_url)
        jd_skills = extract_skills(jd_text)
        doc, doc_lines = read_docx(uploaded_file)
        update_skills_only(doc, doc_lines, jd_skills)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        st.success("✅ Skills updated successfully!")
        st.download_button("⬇️ Download Updated Resume", out, file_name="resume_updated.docx")
